import os
import fitz  # PyMuPDF
from pdf2image import convert_from_path
from paddleocr import PaddleOCR
from docx import Document
import numpy as np
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
import argparse
import logging

# Set logging level
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Initialize PaddleOCR model with improved settings
ocr = PaddleOCR(
    use_textline_orientation=True,  # Use the new parameter instead of use_angle_cls
    lang='ch',  # Support both Chinese and English
    show_log=False,  # Reduce log output
    det_db_thresh=0.3,  # Lower detection threshold for better text detection
    det_db_box_thresh=0.5,  # Lower box threshold
    rec_batch_num=6  # Increase batch size for faster processing
)

def is_valid_content(text, min_valid_ratio=0.1):
    """Check if the text contains enough valid content"""
    total_chars = len(text)
    valid_chars = len(re.findall(r'\S', text))  # non-whitespace characters
    return valid_chars / total_chars > min_valid_ratio if total_chars > 0 else False

def convert_pdf_to_docx_with_ocr(pdf_path, docx_path):
    """Convert PDF to DOCX using OCR with improved text organization"""
    logging.info(f"Processing with OCR: {pdf_path}")
    doc = Document()
    try:
        # Use higher DPI for better OCR quality
        pages = convert_from_path(pdf_path, dpi=400, first_page=1, last_page=None)
        
        for page_num, page in enumerate(pages):
            logging.info(f"OCR processing page {page_num+1}/{len(pages)}")
            image_np = np.array(page.convert('RGB'))
            
            # Run OCR with confidence scores
            result = ocr.ocr(image_np, cls=True)
            
            if result and result[0]:
                # Sort text blocks by vertical position (top to bottom)
                # This helps maintain the reading order
                sorted_blocks = sorted(result[0], key=lambda x: x[0][0][1])  # Sort by y-coordinate of first point
                
                # Group text lines that are likely to be in the same paragraph
                paragraphs = []
                current_paragraph = []
                last_y = None
                
                for block in sorted_blocks:
                    # Extract text and confidence
                    text, confidence = block[1]
                    
                    # Skip low confidence text
                    if confidence < 0.6:  # Arbitrary threshold
                        continue
                        
                    # Get coordinates
                    coords = block[0]
                    y_coord = coords[0][1]  # Top y-coordinate
                    
                    # If this is a new paragraph (based on vertical spacing)
                    if last_y is not None and abs(y_coord - last_y) > 20:  # Arbitrary threshold
                        if current_paragraph:
                            paragraphs.append(" ".join(current_paragraph))
                            current_paragraph = []
                    
                    current_paragraph.append(text)
                    last_y = coords[2][1]  # Bottom y-coordinate
                
                # Add the last paragraph
                if current_paragraph:
                    paragraphs.append(" ".join(current_paragraph))
                
                # Add paragraphs to document
                for paragraph_text in paragraphs:
                    doc.add_paragraph(paragraph_text)
            else:
                # If OCR failed to detect any text
                doc.add_paragraph(f"[OCR could not extract text from page {page_num+1}]")
            
            # Add page break after each page except the last one
            if page_num < len(pages) - 1:
                doc.add_page_break()
        
        doc.save(docx_path)
        logging.info(f"OCR processing completed and saved to {docx_path}")
    except Exception as e:
        logging.error(f"Error processing PDF with OCR: {pdf_path}")
        logging.error(f"Error details: {str(e)}")
        # Create a document with error information
        error_doc = Document()
        error_doc.add_paragraph(f"Error processing PDF with OCR: {str(e)}")
        error_doc.save(docx_path)

def convert_pdf_to_docx(pdf_path, docx_path):
    """Convert PDF to DOCX, use OCR if MuPDF error occurs or text extraction is poor"""
    try:
        doc = Document()
        with fitz.open(pdf_path) as pdf_document:
            for page_num, page in enumerate(pdf_document):
                # Try different text extraction methods
                text = page.get_text("text")
                
                # If text extraction yields poor results, try blocks mode
                if not text.strip() or len(text) < 50:  # Arbitrary threshold
                    blocks = page.get_text("blocks")
                    if blocks:
                        text_blocks = [b[4] for b in blocks if len(b) > 4]
                        text = "\n\n".join(text_blocks)
                
                # If still poor results, try dict mode which preserves more formatting
                if not text.strip() or len(text) < 50:
                    dict_text = page.get_text("dict")
                    if dict_text and "blocks" in dict_text:
                        text_parts = []
                        for block in dict_text["blocks"]:
                            if "lines" in block:
                                for line in block["lines"]:
                                    if "spans" in line:
                                        line_text = " ".join(span["text"] for span in line["spans"] if "text" in span)
                                        if line_text.strip():
                                            text_parts.append(line_text)
                        text = "\n".join(text_parts)
                
                # If text is still not good, mark for OCR processing later
                if not is_valid_content(text):
                    logging.warning(f"Poor text extraction on page {page_num+1}, may need OCR")
                
                # Add the text to the document
                doc.add_paragraph(text)
                doc.add_page_break()
        
        doc.save(docx_path)
        
        # Check if the document has valid content overall
        with open(docx_path, "rb") as f:
            doc = Document(f)
            total_text = " ".join([paragraph.text for paragraph in doc.paragraphs])
            if not is_valid_content(total_text, min_valid_ratio=0.2):
                logging.warning(f"Poor overall text extraction, switching to OCR: {pdf_path}")
                convert_pdf_to_docx_with_ocr(pdf_path, docx_path)
                
    except fitz.fitz.FileDataError as e:
        logging.warning(f"MuPDF error, switching to OCR: {pdf_path}")
        logging.warning(f"Error details: {str(e)}")
        convert_pdf_to_docx_with_ocr(pdf_path, docx_path)
    except Exception as e:
        logging.error(f"Error processing PDF: {pdf_path}")
        logging.error(f"Error details: {str(e)}")

def process_file(file_path, output_folder, output_filename=None):
    """Process a single file
    
    Args:
        file_path: Path to the PDF file
        output_folder: Folder to save the output file
        output_filename: Optional specific output filename (without path)
    
    Returns:
        Tuple of (input_path, output_path) or None if processing failed
    """
    filename = os.path.basename(file_path)
    
    if output_filename:
        docx_path = os.path.join(output_folder, output_filename)
    else:
        docx_path = os.path.join(output_folder, os.path.splitext(filename)[0] + '.docx')
    
    if file_path.endswith('.pdf'):
        logging.info(f"Converting PDF: {filename}")
        convert_pdf_to_docx(file_path, docx_path)
        return file_path, docx_path
    else:
        return None

def process_files(input_folder, output_folder, max_workers=4):
    """Process all files in the folder"""
    processed_files = []
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_file = {}
        for filename in os.listdir(input_folder):
            file_path = os.path.join(input_folder, filename)
            if filename.endswith('.pdf'):
                future = executor.submit(process_file, file_path, output_folder, None)
                future_to_file[future] = file_path

        for future in as_completed(future_to_file):
            file_path = future_to_file[future]
            try:
                result = future.result()
                if result:
                    processed_files.append(result)
            except Exception as e:
                logging.error(f"Error processing file: {file_path}")
                logging.error(f"Error details: {str(e)}")
    
    return processed_files

def check_and_fix_files(processed_files, size_threshold):
    """Check and fix files based on size and content"""
    for pdf_path, docx_path in processed_files:
        need_ocr = False
        
        # Check file size
        if os.path.getsize(docx_path) < size_threshold:
            logging.warning(f"Found file smaller than {size_threshold/1024}KB: {docx_path}")
            need_ocr = True
        
        # Check file content
        if not need_ocr:
            doc = Document(docx_path)
            total_text = " ".join([paragraph.text for paragraph in doc.paragraphs])
            if not is_valid_content(total_text):
                logging.warning(f"File {docx_path} contains mainly blank content.")
                need_ocr = True
        
        # Reprocess with OCR if needed
        if need_ocr:
            logging.info(f"Reprocessing PDF with OCR: {pdf_path}")
            convert_pdf_to_docx_with_ocr(pdf_path, docx_path)
            logging.info(f"Reprocessed and saved: {docx_path}")

def process_single_file(input_file, output_file, size_threshold):
    """Process a single PDF file and validate the result
    
    Args:
        input_file: Path to the input PDF file
        output_file: Path to the output DOCX file
        size_threshold: Size threshold in bytes for quality check
        
    Returns:
        bool: True if processing was successful, False otherwise
    """
    logging.info(f"Processing single file: {input_file}")
    
    # Create output directory if it doesn't exist
    output_dir = os.path.dirname(output_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Process the file with the specific output filename
    output_filename = os.path.basename(output_file)
    result = process_file(input_file, output_dir, output_filename)
    if not result:
        logging.error(f"Failed to process file: {input_file}")
        return False
    
    # Get the paths
    pdf_path, docx_path = result
    
    # Check and fix the file if needed
    need_ocr = False
    
    # Check file size
    if os.path.getsize(docx_path) < size_threshold:
        logging.warning(f"Found file smaller than {size_threshold/1024}KB: {docx_path}")
        need_ocr = True
    
    # Check file content
    if not need_ocr:
        doc = Document(docx_path)
        total_text = " ".join([paragraph.text for paragraph in doc.paragraphs])
        if not is_valid_content(total_text):
            logging.warning(f"File {docx_path} contains mainly blank content.")
            need_ocr = True
    
    # Reprocess with OCR if needed
    if need_ocr:
        logging.info(f"Reprocessing PDF with OCR: {input_file}")
        convert_pdf_to_docx_with_ocr(input_file, docx_path)
        logging.info(f"Reprocessed and saved: {docx_path}")
    
    return True

def main():
    parser = argparse.ArgumentParser(description="Convert PDF files to DOCX format")
    # Define two mutually exclusive groups for input/output
    file_group = parser.add_argument_group('Single file processing')
    file_group.add_argument("--input_file", help="Input PDF file path")
    file_group.add_argument("--output_file", help="Output DOCX file path")
    
    folder_group = parser.add_argument_group('Batch processing')
    folder_group.add_argument("--input_folder", help="Input folder path containing PDF files")
    folder_group.add_argument("--output_folder", help="Output folder path for DOCX files")
    
    # Common arguments
    parser.add_argument("--max_workers", type=int, default=4, help="Maximum number of worker threads (for batch processing)")
    parser.add_argument("--size_threshold", type=int, default=15, help="Small file threshold (KB)")
    args = parser.parse_args()
    
    # Convert size threshold from KB to bytes
    size_threshold = args.size_threshold * 1024

    # Check if we're processing a single file or a folder
    if args.input_file and args.output_file:
        # Single file processing
        if not os.path.exists(args.input_file):
            logging.error(f"Input file does not exist: {args.input_file}")
            return
        
        if not args.input_file.endswith('.pdf'):
            logging.error(f"Input file is not a PDF: {args.input_file}")
            return
            
        success = process_single_file(args.input_file, args.output_file, size_threshold)
        if success:
            logging.info(f"File processed successfully: {args.output_file}")
        else:
            logging.error(f"Failed to process file: {args.input_file}")
    
    elif args.input_folder and args.output_folder:
        # Batch processing (original functionality)
        if not os.path.exists(args.output_folder):
            os.makedirs(args.output_folder)

        processed_files = process_files(args.input_folder, args.output_folder, args.max_workers)
        logging.info("Checking and fixing files...")
        check_and_fix_files(processed_files, size_threshold)
        logging.info("All files processed.")
    
    else:
        logging.error("You must specify either --input_file and --output_file for single file processing OR --input_folder and --output_folder for batch processing")
        parser.print_help()

if __name__ == "__main__":
    main()